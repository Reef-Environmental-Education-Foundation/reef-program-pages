"""
generate_booking_package.py -- the live-Airtable version of the two
proof-of-concept scripts built 2026-09-03 (gen_contract.py and
generate_confirmed_page.py). Given a Bookings record ID, this pulls fresh
data straight from the Airtable REST API (no hardcoded dict) and produces:

  1. bookings/<slug>/data.js   -- the confirmed/pre-trip page's data,
     matching the schema in assets/render.js.
  2. contracts/<slug>-contract.docx -- the merged contract document.

This is the script .github/workflows/publish-booking.yml calls. It
requires an AIRTABLE_API_KEY environment variable (a read-only personal
access token scoped to the two bases below) -- see HANDOFF_INSTRUCTIONS.md
for how to create one and add it as a GitHub Actions secret.

USAGE:
    AIRTABLE_API_KEY=... python generate_booking_package.py <booking_record_id>

Honesty note on what this script can and cannot do (unchanged from the
two POCs it replaces): Airtable has no field recording which day/time slot
a requested Activity falls into for a given booking, so the day-by-day
itinerary content is NOT derived from Activities Requested here. Instead
this script looks up a per-booking itinerary JSON under itineraries/
(captured from that booking's already-sent, human-approved proposal) and
reshapes it into the confirmed page's schema. If a booking has no file
under itineraries/<record_id>.json yet, this script fails loudly rather
than fabricating an itinerary -- see NoItineraryCaptured below. Building a
real day/time data model so this step isn't needed is tracked separately
(OXP_System_Connection_Roadmap_2026-09-03.md, priority #4) and is out of
scope here per Martha's 2026-09-03 direction to skip the intake/data-model
work for this push.
"""

import copy
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import requests
import docx

AIRTABLE_API_KEY = os.environ.get("AIRTABLE_API_KEY")
API_ROOT = "https://api.airtable.com/v0"

BOOKINGS_BASE = "app3FIECuG8iQr7kY"
BOOKINGS_TABLE = "Bookings"
EPO_BASE = "appCK2qlT3WHdhYKd"
PERSONNEL_TABLE = "REEF Personnel"

REPO_ROOT = Path(__file__).parent
TEMPLATE_PATH = REPO_ROOT / "templates" / "Benjamin_School_Expedition_Contract_TEMPLATE.docx"
ITINERARIES_DIR = REPO_ROOT / "itineraries"
BOOKINGS_OUT_DIR = REPO_ROOT.parent / "bookings"
CONTRACTS_OUT_DIR = REPO_ROOT.parent / "contracts"


class NoItineraryCaptured(Exception):
    """Raised when a booking has no captured day-by-day content yet.
    See the module docstring -- this is deliberate, not a bug to patch
    around with fabricated content."""


def airtable_get(base_id, table_name, record_id):
    if not AIRTABLE_API_KEY:
        raise RuntimeError("AIRTABLE_API_KEY environment variable is not set.")
    url = f"{API_ROOT}/{base_id}/{table_name}/{record_id}"
    resp = requests.get(url, headers={"Authorization": f"Bearer {AIRTABLE_API_KEY}"}, timeout=30)
    resp.raise_for_status()
    return resp.json()  # {"id": ..., "createdTime": ..., "fields": {...}}


def slugify(text):
    return re.sub(r"-+", "-", re.sub(r"[^a-z0-9]+", "-", text.lower())).strip("-")


def money(n):
    return f"${n:,.2f}"


def date_pretty(iso):
    return datetime.strptime(iso[:10], "%Y-%m-%d").strftime("%b %-d, %Y")


def us_date(iso):
    d = datetime.strptime(iso[:10], "%Y-%m-%d")
    return f"{d.month}/{d.day}/{d.year % 100:02d}"


def fetch_booking_data(record_id):
    """Pulls the Bookings record + its linked REEF Personnel contact and
    normalizes them into one plain dict. Every value here traces to a
    live Airtable field -- field names match Airtable's REST API exactly
    (see list_tables_for_base output captured 2026-09-03)."""
    rec = airtable_get(BOOKINGS_BASE, BOOKINGS_TABLE, record_id)
    f = rec["fields"]

    personnel_record_id = f.get("REEF Booking Contact (EPO Personnel Record ID)")
    reef_contact = {}
    if personnel_record_id:
        p = airtable_get(EPO_BASE, PERSONNEL_TABLE, personnel_record_id)["fields"]
        reef_contact = {
            "name": p.get("Full name", ""),
            "title": p.get("Customer-Facing Title", ""),
            "email": p.get("Email", ""),
            "phone": p.get("Phone", ""),
            "welcome_line": p.get("Proposal Welcome Line", ""),
        }

    return {
        "record_id": record_id,
        "org_name": f.get("Organization Name", ""),
        "contact_name": f.get("Primary Contact Name", ""),
        "contact_role": f.get("Contact Role", ""),
        "contact_email": f.get("Contact Email", ""),
        "contact_phone": f.get("Contact Phone", ""),
        "arrival_date": f.get("Arrival Date"),
        "departure_date": f.get("Departure Date"),
        "location": f.get("Location", ""),
        "students": f.get("# Student Participants", 0),
        "chaperones": f.get("# Total Chaperones", 0),
        "booking_type": f.get("Booking Type", ""),
        "total_package_price": f.get("Total Package Price", 0),
        "price_per_paid_space": f.get("Price Per Paid Space", 0),
        "status": f.get("Status", ""),
        "proposal_response": f.get("Proposal Response", ""),
        "agreement_contract_status": f.get("Agreement/Contract Status"),
        "deposit_payment_status": f.get("Deposit/Payment Status"),
        "payment_tier": f.get("Payment Tier", ""),
        "deposit_due_now": f.get("Deposit Due Now", 0),
        "payment2_amount": f.get("Payment 2 Amount", 0),
        "payment2_due_date": f.get("Payment 2 Due Date"),
        "final_amount": f.get("Final Payment Amount", 0),
        "final_due_date": f.get("Final Payment Due Date"),
        "reef_contact": reef_contact,
        "reef_program_fees_total": f.get("REEF Program Fees Total", 0),
        "reef_supplies_total": f.get("REEF Supplies Total", 0),
        "campus_facility_fee_total": f.get("Campus Facility Fee Total", 0),
        "proposal_included": f.get("Proposal — What's Included", ""),
        "proposal_not_included": f.get("Proposal — What's Not Included", ""),
        "proposal_assumptions": f.get("Proposal — Assumptions", ""),
    }


def load_itinerary(record_id):
    path = ITINERARIES_DIR / f"{record_id}.json"
    if not path.exists():
        raise NoItineraryCaptured(
            f"No captured itinerary for booking {record_id} at {path}. "
            "This script deliberately does not fabricate a day-by-day schedule -- "
            "export the itinerary JSON from that booking's already-sent proposal first "
            "(see OXP_System_Connection_Roadmap_2026-09-03.md, priority #4, for the real fix)."
        )
    with open(path) as fh:
        return json.load(fh)


# ---------------------------------------------------------------- Confirmed page

def build_confirmed_page_data(b):
    days_raw = load_itinerary(b["record_id"])
    days = []
    for d in days_raw:
        days.append({
            "dayNumber": d["dayNumber"],
            "totalDays": d["totalDays"],
            "title": d["title"],
            "theme": d["theme"],
            "morningLabel": d["blocks"][0]["title"] if d["blocks"] else "",
            "afternoonLabel": d["blocks"][-1]["title"] if len(d["blocks"]) > 1 else "",
            "learningOutcome": "; ".join(s.split(" (")[0] for s in d.get("studentsWill", [])[:2]),
            "blocks": d["blocks"],
            "studentsWill": d.get("studentsWill", []),
            "outcomesNote": d.get("outcomesNote"),
        })

    status = b["agreement_contract_status"]
    if status == "Signed":
        action_needed = {"show": False}
        agreement = {"status": "Signed", "zohoSignUrl": None, "lastUpdated": None}
    elif status == "Sent":
        action_needed = {
            "show": True,
            "headline": "Your agreement is ready to sign",
            "detail": "Review and sign your Ocean Explorers agreement to lock in your dates.",
            "ctaText": "Review & Sign Agreement",
            # Zoho Sign integration is out of scope for this build (Martha, 2026-09-03).
            "ctaUrl": None,
        }
        agreement = {"status": "Sent", "zohoSignUrl": None, "lastUpdated": None}
    else:
        action_needed = {
            "show": True,
            "headline": "Your contract is being prepared",
            "detail": "REEF is finalizing your agreement based on the proposal you approved. You'll receive it here once it's ready to sign.",
            "ctaText": None,
            "ctaUrl": None,
        }
        agreement = {"status": "Not Sent", "zohoSignUrl": None, "lastUpdated": None}

    payment_schedule = {
        "tier": b["payment_tier"],
        "items": [
            {"label": "Deposit (due now)", "amount": money(b["deposit_due_now"]), "dueDate": None},
            {"label": "Payment 2", "amount": money(b["payment2_amount"]),
             "dueDate": date_pretty(b["payment2_due_date"]) if b["payment2_due_date"] else None},
            {"label": "Final Payment", "amount": money(b["final_amount"]),
             "dueDate": date_pretty(b["final_due_date"]) if b["final_due_date"] else None},
        ],
        "total": money(b["total_package_price"]),
        "note": "Generated from REEF Bookings' payment-schedule fields (added 2026-09-03); "
                "confirm the 50/50 Payment 2 / Final split assumption before treating as final.",
    }

    return {
        "docType": "pretrip",
        "assetDepth": 2,
        "meta": {
            "sampleFlag": False,
            "generatedFrom": b["record_id"],
            "generatedNote": f"Generated by generate_booking_package.py from live Airtable "
                              f"fields on {datetime.utcnow().strftime('%Y-%m-%d')}.",
        },
        "program": {
            "name": "Florida Keys Marine Science Expedition",
            "track": "Expedition",
            "groupName": b["org_name"],
            "schoolOrg": b["org_name"],
            "gradeLevel": b["contact_role"],
            "groupSize": f"{b['students']} Students / {b['chaperones']} Chaperones",
            "location": "REEF Campus, Key Largo",
            "dates": {
                "label": "Confirmed" if b["status"] == "Confirmed" else "Proposed",
                "range": f"{date_pretty(b['arrival_date'])} - {date_pretty(b['departure_date'])}",
            },
        },
        "contacts": {
            "educatorName": b["contact_name"],
            "reefEducatorName": b["reef_contact"].get("name", ""),
            "reefEducatorTitle": b["reef_contact"].get("title", ""),
            "reefEducatorWelcomeLine": b["reef_contact"].get("welcome_line", ""),
            "reefPhone": b["reef_contact"].get("phone", ""),
            "reefEmail": b["reef_contact"].get("email", ""),
        },
        "hero": {
            "kicker": "OCEAN EXPLORERS\nEXPEDITION PACKET",
            "eyebrowTag": "Florida Keys Marine Science Expedition",
            "headline": "From Student to Scientist in Key Largo",
            "promise": "Turn the ocean into your classroom. Your students won't just study marine science -- they become part of it.",
            "imageUrl": "../../assets/photos/hero-reef-shark.jpg",
            "imageCredit": "Photo: Jeffrey Haines / REEF",
        },
        "actionNeeded": action_needed,
        "agreement": agreement,
        "paymentSchedule": payment_schedule,
        "nextSteps": {
            "items": [
                "Review and sign your agreement once it's sent (see above).",
                "Return your group's signed waivers and health/medical forms.",
                "Confirm final headcount with your REEF educator at least 2 weeks before arrival.",
                "Reach out any time with questions before your expedition.",
            ],
        },
        "welcome": {
            "body": [
                "We're glad your group is joining us. This packet lays out what to expect from your "
                "Florida Keys Marine Science Expedition -- where your students will identify reef fish, "
                "explore real coral reef, mangrove, and seagrass habitats, and practice the same "
                "citizen-science methods REEF's volunteer network uses across the Caribbean and beyond.",
                "This is not a sightseeing trip. It's a working expedition: your students will observe, "
                "identify, survey, investigate, and contribute -- and leave with a real sense of what it "
                "means to practice marine science, not just read about it.",
            ],
            "signOff": "The REEF Ocean Explorers Team",
        },
        "glanceNote": "A quick-scan summary for planning. Full detail -- including “students will” "
                      "outcomes and gear notes -- follows on the day-by-day pages.",
        "days": days,
    }


# ---------------------------------------------------------------- Contract

def set_paragraph_text(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def replace_exact(doc, old, new, occurrence=0):
    count = 0
    for p in doc.paragraphs:
        if p.text.strip() == old:
            if count == occurrence:
                set_paragraph_text(p, new)
                return True
            count += 1
    return False


def insert_watermark_banner(doc):
    banner = doc.paragraphs[0].insert_paragraph_before()
    run = banner.add_run(
        "⚠ GENERATED FROM LIVE AIRTABLE DATA -- REVIEW BEFORE SENDING."
    )
    run.bold = True
    run.font.size = docx.shared.Pt(12)
    run.font.color.rgb = docx.shared.RGBColor(0xB2, 0x1D, 0x1D)


def build_contract(b, out_path):
    """Same merge logic proven in gen_contract.py 2026-09-03, now driven
    by live-fetched values instead of a hardcoded dict. This does not
    re-derive the freeform program description / topics text (those were
    hand-composed for the Riverside Academy proof-of-concept and Airtable
    has no single field holding "topics covered" as a clean list yet) --
    those two paragraphs are left as the template's originals for a human
    to review and edit before sending, same as today's manual process,
    until that content gets a real field to live in."""
    doc = docx.Document(TEMPLATE_PATH)
    insert_watermark_banner(doc)

    replace_exact(doc, "The Benjamin School", b["org_name"], occurrence=0)
    replace_exact(doc, "The Benjamin School", b["org_name"], occurrence=0)
    replace_exact(doc, "Erin Gigele", b["contact_name"])
    replace_exact(doc, "7th Grade Marine Science", b["contact_role"])
    replace_exact(doc, "561-626-3747 (ext. 3362)", b["contact_phone"])
    replace_exact(doc, "erin.ryan@thebenjaminschool.org", b["contact_email"])
    replace_exact(doc, "Contract Date: 9/2/2025", f"Contract Date: {datetime.utcnow().strftime('%-m/%-d/%Y')}")
    replace_exact(doc, "Program Dates: 4/16/26- 4/17/26",
                  f"Program Dates: {us_date(b['arrival_date'])} - {us_date(b['departure_date'])}")
    replace_exact(
        doc,
        "Program Location: Key Largo: REEF Campus, Pennekamp Coral Reef State Park, MOTE Coral Nursey",
        f"Program Location: {b['location']}",
    )
    replace_exact(doc, "Total Participants: 90", f"Total Participants: {b['students'] + b['chaperones']}")
    replace_exact(doc, "Total Program Cost Rate: $24,510",
                  f"Total Program Cost Rate: {money(b['total_package_price'])}")
    replace_exact(
        doc,
        "Program Rate Per Person (90): $272",
        f"Program Rate Per Person ({b['students'] + b['chaperones']}): {money(round(b['price_per_paid_space']))}",
    )
    replace_exact(doc, "$1000 to secure group space", f"{money(b['deposit_due_now'])} to secure group space")
    if b["payment2_due_date"]:
        replace_exact(
            doc, "$11,755 by 8/18/2025 (Half after 1000)",
            f"{money(b['payment2_amount'])} by {date_pretty(b['payment2_due_date'])} (Payment 2)",
        )
    if b["final_due_date"]:
        replace_exact(
            doc, "$11,755 by 1/16/2026 (Final Payment)",
            f"{money(b['final_amount'])} by {date_pretty(b['final_due_date'])} (Final Payment)",
        )
    replace_exact(doc, "Print Name\tRose Kelly", f"Print Name\t{b['reef_contact'].get('name', '')}")

    doc.save(out_path)


BOOKING_PAGE_SHELL_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,400;9..144,500;9..144,600;9..144,700&family=Public+Sans:wght@400;500;600;700&display=swap">
<title>{title}</title>
<link rel="stylesheet" href="../../assets/styles.css">
</head>
<body>
  <div class="page">
    <div id="sample-flag"></div>
    <div id="hero"></div>
    <div id="mini-nav"></div>
    <div id="action-needed"></div>
    <div id="snapshot"></div>
    <div id="welcome"></div>
    <div id="glance"></div>
    <div id="day-by-day"></div>
    <div id="students-will-do"></div>
    <div id="gear"></div>
    <div id="next-steps"></div>
    <div id="closing-cta"></div>
    <div id="site-footer"></div>
  </div>

  <script src="data.js"></script>
  <script src="../../assets/render.js"></script>
</body>
</html>
"""
# ^ This is the same thin shell every booking page in bookings/<slug>/ uses
# (identical to bookings/sample-ocean-explorers-2day/index.html, which is
# hand-maintained separately since it's a fixed sample, not a generated
# booking). Every generated booking gets its own copy of this file because
# the earlier version of this script wrote data.js alone -- that page
# 404'd on GitHub Pages until this was added (2026-09-03, after Martha's
# first live test run). If the shared shell markup ever changes, update it
# both here and in the sample page.


def main():
    if len(sys.argv) != 2:
        print("Usage: python generate_booking_package.py <booking_record_id>", file=sys.stderr)
        sys.exit(1)
    record_id = sys.argv[1]

    b = fetch_booking_data(record_id)
    slug = slugify(b["org_name"])

    BOOKINGS_OUT_DIR.mkdir(parents=True, exist_ok=True)
    CONTRACTS_OUT_DIR.mkdir(parents=True, exist_ok=True)

    booking_dir = BOOKINGS_OUT_DIR / slug
    booking_dir.mkdir(exist_ok=True)
    confirmed_data = build_confirmed_page_data(b)
    with open(booking_dir / "data.js", "w") as f:
        f.write("/* GENERATED by generate_booking_package.py -- do not hand-edit. */\n")
        f.write("window.BOOKING_DATA = " + json.dumps(confirmed_data, indent=2) + ";\n")
    print(f"Wrote {booking_dir / 'data.js'}")

    page_title = f"{b['org_name']} — Expedition Packet"
    with open(booking_dir / "index.html", "w") as f:
        f.write(BOOKING_PAGE_SHELL_TEMPLATE.format(title=page_title))
    print(f"Wrote {booking_dir / 'index.html'}")

    contract_path = CONTRACTS_OUT_DIR / f"{slug}-contract.docx"
    build_contract(b, contract_path)
    print(f"Wrote {contract_path}")


if __name__ == "__main__":
    main()
